# Document correction
# Takes in a document missing a definite article or preposition and makes suggestions for each sentence

# Imports python-docx
DOCPATH = "C:\\Original.docx"

from docx import Document

# Imports Distilbert, TensorFlow, numpy and builds a model for mask fill

import sys
import itertools

import os

import math
from pytorch_pretrained_bert import BertForMaskedLM
import torch
import pandas as pd
import math

from transformers import TFAutoModelForMaskedLM

bertMaskedLM = BertForMaskedLM.from_pretrained('bert-base-uncased').eval()

model_checkpoint = "distilbert-base-uncased"
model = TFAutoModelForMaskedLM.from_pretrained(model_checkpoint)

model(model.dummy_inputs)  # Build the model
model.summary()

from transformers import AutoTokenizer

tokenizer = AutoTokenizer.from_pretrained(model_checkpoint)
import numpy as np
import tensorflow as tf

# Sorts prepositions by their computed sentence score, then returns a list of the top 5
# Parameters: A list of tuples, prep_list
# Returns: A list of tuples
def get_top_5(prep_list):
    prep_list = sorted(prep_list.items(), key=lambda x: x[1])[:5]
    return prep_list

# Computes the cross-entropy loss for a sentence with function words added via mask fill
# Parameters: A string, sentence
# Returns: A float, the sentence score
def get_score(sentence):
    tokenize_input = tokenizer.tokenize(sentence)
    tensor_input = torch.tensor([tokenizer.convert_tokens_to_ids(tokenize_input)])
    predictions=bertMaskedLM(tensor_input)
    loss_fct = torch.nn.CrossEntropyLoss()
    loss = loss_fct(predictions.squeeze(),tensor_input.squeeze()).data 
    return math.exp(loss)

# Mask fill article suggestion function: Checks to see if, for the mask in question, one of the top
# five suggestions is 'the'. If this is the case, the mask is filled and the new sentence is
# appended to the new document.
# Parameters: A string - sentence, a string - article, a document object - target_document
# Returns: A string
def suggest_articles(sentence, article, target_document):
    inputs = tokenizer(sentence, return_tensors="np")
    token_logits = model(**inputs).logits

    mask_token_index = np.argwhere(inputs["input_ids"] == tokenizer.mask_token_id)[0, 1]
    mask_token_logits = token_logits[0, mask_token_index, :]

    top_5_tokens = np.argsort(-mask_token_logits)[:5].tolist()
    top_5_decoded = []
    for token in top_5_tokens:
        top_5_decoded.append(tokenizer.decode([token]))

    #print(tokenizer.decode([token]))
    if(article in top_5_decoded) or f"{article.capitalize()}" in top_5_decoded:
        sentence = sentence.replace(tokenizer.mask_token,f'{article}')
        sentence = sentence.lstrip()
        if(sentence.startswith(f"{article}")):
            sentence = sentence.replace(f"{article}", f"{article.capitalize()}")
        if '[MASK]' not in sentence:
            return sentence
            
# Provides a list of candidate prepositions if the preposition provided is not in the top 5
# suggested by mask fill
# Parameters: A string - sentence, a list - preposition_list, a tokenized list of the words
# in the sentence - tokenized_sentence, the index of the mask token - mask_index, the word
# the mask token replaced - word
# Returns: A string, sentence
def suggest_prepositions(sentence, preposition_list, tokenized_sentence, mask_index, word):
    if "[MASK]" in sentence:
        inputs = tokenizer(sentence, return_tensors="np")
        token_logits = model(**inputs).logits

        mask_token_index = np.argwhere(inputs["input_ids"] == tokenizer.mask_token_id)[0, 1]
        mask_token_logits = token_logits[0, mask_token_index, :]

        top_10_token = np.argsort(-mask_token_logits)[:10].tolist()

        prepositions_to_suggest = []
        for token in top_10_token:
            prepositions_to_suggest.append(tokenizer.decode([token]))
        # print(f"Suggested mask fills: {prepositions_to_suggest}")
        # sentence_list = {}
        prep_list = {}
        for i in prepositions_to_suggest:
            if i in preposition_list and i != tokenized_sentence[mask_index]:
                suggested_sentence = sentence
                suggested_sentence = suggested_sentence.replace(tokenizer.mask_token,f'{i}').strip('\n')
                suggested_sentence = suggested_sentence.lstrip()
                if(suggested_sentence.startswith(f"{i}")):
                    suggested_sentence = suggested_sentence.replace(f"{i}", f"{i.capitalize()}")
                # print(suggested_sentence, get_score(suggested_sentence))
                prep_score = get_score(suggested_sentence)
                prep_list[i] = prep_score
        top_5 = get_top_5(prep_list)
        top_5_list = list(top_5)
        # if word in top_5:
        #     return sentence.replace("[MASK]", word)
        # else:
        sentence = sentence.replace("[MASK]", str(list(prep_list.keys())))
        return sentence

# Inserts a mask before each word in the sentence and runs the mask fill task on the mask token.
# Finally, it adds the sentence to the working document
# Parameters: A string - sentence, a list - articles, a document object - working_document
def article_suggestion(sentence, articles, working_document):
    tokenized_sentence = sentence.split(" ")
    
    # For each token in the sentence, place a mask before it and run the mask fill function
    for word in tokenized_sentence:
        # print(tokenized_sentence.index(word))
        masked_sentence = tokenized_sentence[:]
        if(masked_sentence.index(word) == 0):
            masked_sentence.insert(0, "[MASK]")
            # print("Sentence with mask: ", *masked_sentence)
            sentence_for_analysis = ''
            for word in masked_sentence:
                sentence_for_analysis = sentence_for_analysis + ' ' + word
            # print(f"Sentence for analysis: {sentence_for_analysis}")
            for article in articles:
                if suggest_articles(sentence_for_analysis, article, working_document):
                    tokenized_sentence = suggest_articles(sentence_for_analysis, article, working_document).split()


        else:
            #masked_sentence[(masked_sentence.index(word)-1)] = tokenized_sentence[tokenized_sentence.index(word)-1]
            masked_sentence.insert(masked_sentence.index(word), "[MASK]")
            # print("Sentence with mask: ", *masked_sentence)
            sentence_for_analysis = ''
            for word in masked_sentence:
                sentence_for_analysis = sentence_for_analysis + ' ' + word
            # print(f"Sentence for analysis: {sentence_for_analysis}")
            for article in articles:
                if suggest_articles(sentence_for_analysis, article, working_document):
                    tokenized_sentence = suggest_articles(sentence_for_analysis, article, working_document).split()

    working_document.add_paragraph(' '.join(tokenized_sentence))

# Replaces each preposition in the sentence with a mask token if it is in the preposition_list. Then runs the
# mask fill function. Finally, returns the sentence
# Parameters: A string, sentence
# Returns: A string, sentence_to_return
def preposition_suggestion(sentence):
    mask_index = -1
    preposition_list = ["aboard", "about", "above", "across", "after", "against", "along", "amid", "among", "anti", "around", "as", "at", "before", "behind", "below", "beneath", "beside", "besides", "between", "beyond", "but", "by", "concerning", "considering", "despite", "down", "during", "except", "excepting", "excluding", "following", "for", "from", "in", "inside", "into", "like", "minus", "near", "of", "off", "on", "onto", "opposite", "out", "outside", "over", "past", "per", "plus", "regarding", "round", "save", "since", "than", "through", "to", "toward", "towards", "under", "underneath", "unlike", "until", "up", "upon", "versus", "via", "with", "within", "without"]
    tokenized_sentence = sentence.split(" ")
    # For each preposition in the sentence, replace it with a mask and run the mask fill function
    for word in tokenized_sentence:
        # print(tokenized_sentence.index(word))
        masked_sentence = tokenized_sentence[:]
        sentence_for_analysis = ''
        #masked_sentence[(masked_sentence.index(word)-1)] = tokenized_sentence[tokenized_sentence.index(word)-1]
        if masked_sentence[masked_sentence.index(word)].lower() in preposition_list:
            masked_sentence[masked_sentence.index(word)] = "[MASK]"
            mask_index = masked_sentence.index("[MASK]")
            # print("Sentence with preposition removed: ", *masked_sentence)
        
        sentence_for_analysis = ''
        for word in masked_sentence:
            sentence_for_analysis = sentence_for_analysis + ' ' + word
        # print(f"Sentence for analysis: {sentence_for_analysis}")
        sentence_to_return = suggest_prepositions(sentence_for_analysis, preposition_list, tokenized_sentence, mask_index, word)
        if sentence_to_return:
            return sentence_to_return
    
# Compares the original document with a generated one using Microsoft Word. Generates a document
# with revisions for the user.
def compare_docs():
    import win32com.client

    path = "C:\\"
    # note the \\ at the end of the path name to prevent a SyntaxError

    #Create the Application word
    Application=win32com.client.gencache.EnsureDispatch("Word.Application")

    # Compare documents
    Application.CompareDocuments(Application.Documents.Open(os.path.join(path, "Original.docx")),
                                Application.Documents.Open(os.path.join(path, "Revised.docx")))

    # Save the comparison document as "Comparison.docx"
    Application.ActiveDocument.SaveAs (FileName = path + "Comparison.docx")
    # Don't forget to quit your Application
    Application.Quit()

def main():
    target_document = Document()
    source_document = Document(DOCPATH)
    source_text = []
    for para in source_document.paragraphs:
        source_text.append(para.text)
    for para in source_text:
        article_suggestion(para, ['the'], target_document)
    target_document.save('C:\\Revised.docx')
    compare_docs()
    

if __name__ == "__main__":
    main()
